"""DOCX file parser for RAG.

Handles extraction of text from Microsoft Word .docx files.
For old .doc files, delegates to DOCParser.
"""

import logging
from pathlib import Path

from rag_module.file_processing.text_cleaner import TextCleaner
from rag_module.file_processing.doc_parser import DOCParser

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DOCXParser:
    """Парсер для .docx файлов.
    
    Для .doc файлов делегирует существующему DOCParser.
    """
    
    def __init__(self) -> None:
        self.text_cleaner = TextCleaner()
        self.doc_parser = DOCParser()
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from .docx or .doc file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid DOCX
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        # For .doc files, use existing DOCParser
        if suffix == '.doc':
            logger.info(f"Delegating .doc file to DOCParser: {file_path.name}")
            return self.doc_parser.extract_text(file_path)
        
        # For .docx files, use python-docx
        logger.info(f"Starting extraction from {file_path.name} ({file_path.stat().st_size} bytes)")
        
        if Document is None:
            raise ValueError("python-docx not installed")
        
        try:
            logger.info(f"Using python-docx for {file_path.name}")
            doc = Document(file_path)
            extracted_text = []
            
            # Extract paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    extracted_text.append(text)
                    paragraph_count += 1
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                row_text.append(para.text)
                    if row_text:
                        extracted_text.append(" | ".join(row_text))
                        table_count += 1
            
            result = "\n".join(extracted_text)
            if result.strip():
                logger.info(f"✓ Successfully extracted {len(result)} chars from {file_path.name} "
                          f"({paragraph_count} paragraphs, {table_count} tables)")
                return result
            else:
                raise ValueError(f"No text extracted from {file_path.name}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path.name}: {type(e).__name__}: {str(e)[:100]}")
            raise ValueError(f"Cannot extract text from {file_path.name}") from e
    
    def get_metadata(self, file_path: Path) -> dict:  # type: ignore
        """Extract DOCX metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            dict: Document metadata
        """
        if file_path.suffix.lower() != '.docx':
            return {}
        
        try:
            doc = Document(file_path)
            props = doc.core_properties
            return {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "created": props.created,
                "modified": props.modified,
            }
        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")
            return {}
