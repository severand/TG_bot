"""Conversation mode handlers for interactive document analysis.

Allows users to:
1. Upload a document
2. Get confirmation that it's ready
3. Send text instructions for analysis
4. Get analysis results based on custom prompts
"""

import logging
import uuid
from pathlib import Path

from aiogram import Router, F
from aiogram.types import Message, Document, File, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.context import FSMContext
from aiogram.utils.keyboard import InlineKeyboardBuilder

from app.config import get_settings
from app.states.conversation import ConversationStates
from app.services.file_processing.converter import FileConverter
from app.services.llm.llm_factory import LLMFactory
from app.services.prompts.prompt_manager import PromptManager
from app.utils.text_splitter import TextSplitter
from app.utils.cleanup import CleanupManager

logger = logging.getLogger(__name__)

router = Router()
config = get_settings()
prompt_manager = PromptManager()
llm_factory = LLMFactory(
    primary_provider=config.LLM_PROVIDER,
    openai_api_key=config.OPENAI_API_KEY or None,
    openai_model=config.OPENAI_MODEL,
    replicate_api_token=config.REPLICATE_API_TOKEN or None,
    replicate_model=config.REPLICATE_MODEL,
)


def get_document_menu() -> InlineKeyboardMarkup:
    """Menu when document is loaded."""
    builder = InlineKeyboardBuilder()
    builder.button(text="🗑️ Clear Document", callback_data="doc_clear")
    builder.button(text="📌 View Document Info", callback_data="doc_info")
    return builder.as_markup()


def get_start_menu() -> InlineKeyboardMarkup:
    """Main start menu with action buttons."""
    builder = InlineKeyboardBuilder()
    builder.button(text="📄 Upload Document", callback_data="mode_upload")
    builder.button(text="📌 View Prompts", callback_data="mode_prompts")
    return builder.as_markup()


@router.message(F.text == "/analyze")
async def cmd_analyze(message: Message, state: FSMContext) -> None:
    """Start interactive document analysis mode.
    
    /analyze - Активирует режим анализа документов
    """
    await state.set_state(ConversationStates.ready)
    
    text = (
        "💬 *Document Analysis Mode*\n\n"
        "Welcome to interactive document analysis!\n\n"
        "*How it works:*\n"
        "1️⃣ Upload a document (PDF, DOCX, TXT, or ZIP)\n"
        "2️⃣ Bot confirms it's ready\n"
        "3️⃣ Send text messages with analysis instructions\n"
        "4️⃣ Bot analyzes based on your commands\n\n"
        "*Example:*\n"
        "You: [Upload contract.pdf]\n"
        "Bot: Document ready! ✅\n"
        "You: 'Identify all risks in this contract'\n"
        "Bot: [Detailed risk analysis]\n\n"
        "Ready? Upload your document or choose an action below."
    )
    
    await message.answer(
        text,
        parse_mode="Markdown",
        reply_markup=get_start_menu(),
    )
    logger.info(f"User {message.from_user.id} started analysis mode")


@router.message(ConversationStates.ready, F.document)
async def handle_document_upload(message: Message, state: FSMContext) -> None:
    """Handle document upload in conversation mode.
    
    Saves document to memory and confirms it's ready.
    """
    if not message.document:
        await message.answer("❌ No document found")
        return
    
    document: Document = message.document
    file_size = document.file_size or 0
    
    # Validate file size
    if file_size > config.MAX_FILE_SIZE:
        max_size_mb = config.MAX_FILE_SIZE / (1024 * 1024)
        await message.answer(
            f"⚠️ File is too large: {file_size / (1024 * 1024):.1f} MB\n"
            f"Maximum allowed: {max_size_mb:.1f} MB"
        )
        return
    
    # Show processing
    status_msg = await message.answer(
        "🔍 Processing document...\n"
        "Downloading and extracting content..."
    )
    
    temp_user_dir = None
    temp_file_path = None
    
    try:
        # Create temp directory
        temp_base = Path(config.TEMP_DIR)
        temp_base.mkdir(exist_ok=True)
        temp_user_dir = CleanupManager.create_temp_directory(
            temp_base,
            message.from_user.id,
        )
        
        # Download file
        bot = message.bot
        file: File = await bot.get_file(document.file_id)
        
        if not file.file_path:
            await message.answer("❌ Failed to get file path")
            await status_msg.delete()
            return
        
        # Generate unique filename
        file_ext = Path(document.file_name or "document").suffix or ".bin"
        temp_file_path = temp_user_dir / f"{uuid.uuid4()}{file_ext}"
        
        await bot.download_file(file.file_path, temp_file_path)
        logger.info(f"Downloaded: {temp_file_path} ({file_size} bytes)")
        
        # Extract text
        await status_msg.edit_text(
            "🔍 Processing document...\n"
            "Extracting text content..."
        )
        
        converter = FileConverter()
        extracted_text = converter.extract_text(temp_file_path, temp_user_dir)
        
        if not extracted_text or not extracted_text.strip():
            await message.answer(
                "⚠️ No text content found in document.\n"
                "Try another file."
            )
            await status_msg.delete()
            return
        
        # Save to state
        await state.update_data(
            document_text=extracted_text,
            document_name=document.file_name or "document",
            document_size=len(extracted_text),
            user_id=message.from_user.id,
        )
        
        # Move to conversation state
        await state.set_state(ConversationStates.waiting_for_command)
        
        # Confirm
        await status_msg.delete()
        
        text = (
            f"✅ *Document Ready!*\n\n"
            f"*File:* `{document.file_name or 'document'}`\n"
            f"*Size:* {len(extracted_text):,} characters\n\n"
            f"Now send me instructions for analysis:\n\n"
            f"*Examples:*\n"
            f"📝 'Summarize this document'\n"
            f"📝 'Identify all risks'\n"
            f"📝 'Extract key points'\n"
            f"📝 'Analyze legal implications'\n\n"
            f"Or use a custom prompt with /prompts"
        )
        
        await message.answer(
            text,
            parse_mode="Markdown",
            reply_markup=get_document_menu(),
        )
        
        logger.info(
            f"Document loaded for user {message.from_user.id}: "
            f"{len(extracted_text)} chars"
        )
    
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        await message.answer(f"❌ Error: {str(e)}")
        await status_msg.delete()
    
    finally:
        # Don't cleanup yet - document is in state
        pass


@router.message(ConversationStates.waiting_for_command, F.text)
async def handle_analysis_command(message: Message, state: FSMContext) -> None:
    """Handle text commands for document analysis.
    
    User sends instructions, bot analyzes document accordingly.
    """
    command = message.text.strip()
    
    if not command:
        await message.answer("Please provide an analysis instruction.")
        return
    
    # Check for special commands
    if command.startswith("/"):
        # Don't handle commands here
        return
    
    data = await state.get_data()
    document_text = data.get("document_text")
    document_name = data.get("document_name")
    
    if not document_text:
        await message.answer(
            "⚠️ No document loaded.\n"
            "Please upload a document first."
        )
        await state.set_state(ConversationStates.ready)
        return
    
    # Set processing state
    await state.set_state(ConversationStates.processing)
    
    # Show processing indicator
    status_msg = await message.answer(
        f"🤖 Analyzing '{document_name}'...\n"
        f"📍 Instruction: {command[:50]}..."
    )
    
    try:
        # Get user's default prompt or use default
        default_prompt_name = data.get("default_prompt", "default")
        prompt = prompt_manager.get_prompt(message.from_user.id, default_prompt_name)
        
        if not prompt:
            prompt = prompt_manager.get_prompt(message.from_user.id, "default")
        
        # Analyze with user's command
        analysis_result = await llm_factory.analyze_document(
            document_text,
            command,  # Use user's instruction as the prompt
            system_prompt=prompt.system_prompt if prompt else None,
            use_streaming=False,
        )
        
        if not analysis_result:
            await message.answer("❌ Analysis failed. Try again.")
            await status_msg.delete()
            await state.set_state(ConversationStates.waiting_for_command)
            return
        
        # Delete processing message
        await status_msg.delete()
        
        # Send result
        splitter = TextSplitter()
        message_count = splitter.count_messages(analysis_result)
        
        if message_count <= 3:
            # Send as text messages
            chunks = splitter.split(analysis_result)
            for i, chunk in enumerate(chunks, 1):
                prefix = f"*[Part {i}/{len(chunks)}]*\n\n" if len(chunks) > 1 else ""
                await message.answer(
                    f"{prefix}{chunk}",
                    parse_mode="Markdown",
                )
        else:
            # Send as file
            from docx import Document
            
            doc = Document()
            doc.add_heading(
                f"Analysis: {document_name}",
                level=1,
            )
            doc.add_paragraph(f"*Instruction:* {command}")
            doc.add_paragraph(analysis_result)
            
            temp_base = Path(config.TEMP_DIR)
            output_file = temp_base / f"analysis_{uuid.uuid4()}.docx"
            doc.save(output_file)
            
            await message.answer_document(
                open(output_file, "rb"),
                caption=f"📄 Analysis result ({message_count} messages)",
            )
        
        logger.info(
            f"Analysis completed for user {message.from_user.id}: "
            f"instruction='{command[:30]}...'"
        )
    
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        await message.answer(f"❌ Error: {str(e)}")
        await status_msg.delete()
    
    finally:
        # Return to waiting for command
        await state.set_state(ConversationStates.waiting_for_command)


@router.callback_query(F.data == "doc_clear")
async def cb_doc_clear(query, state: FSMContext) -> None:
    """Clear loaded document and return to ready."""
    await state.clear()
    await state.set_state(ConversationStates.ready)
    
    text = (
        "🗑️ Document cleared!\n\n"
        "Ready for a new document. Upload one to get started."
    )
    
    await query.message.edit_text(
        text,
        parse_mode="Markdown",
        reply_markup=get_start_menu(),
    )
    await query.answer()
    logger.info(f"User {query.from_user.id} cleared document")


@router.callback_query(F.data == "doc_info")
async def cb_doc_info(query, state: FSMContext) -> None:
    """Show loaded document information."""
    data = await state.get_data()
    document_name = data.get("document_name", "Unknown")
    document_size = data.get("document_size", 0)
    
    text = (
        f"📌 *Document Information*\n\n"
        f"*Name:* `{document_name}`\n"
        f"*Size:* {document_size:,} characters\n\n"
        f"Continue sending analysis instructions."
    )
    
    await query.message.edit_text(
        text,
        parse_mode="Markdown",
        reply_markup=get_document_menu(),
    )
    await query.answer()


@router.callback_query(F.data == "mode_upload")
async def cb_mode_upload(query, state: FSMContext) -> None:
    """Ready to receive upload."""
    await state.set_state(ConversationStates.ready)
    
    text = (
        "📄 *Ready for Upload*\n\n"
        "Send me a document file:\n"
        "• PDF\n"
        "• DOCX (Word)\n"
        "• TXT\n"
        "• ZIP (with multiple files)\n\n"
        "Maximum size: 20 MB"
    )
    
    await query.message.edit_text(
        text,
        parse_mode="Markdown",
    )
    await query.answer()


@router.callback_query(F.data == "mode_prompts")
async def cb_mode_prompts(query, state: FSMContext) -> None:
    """Show prompts menu."""
    await query.answer("Use /prompts to manage your analysis prompts", show_alert=True)
