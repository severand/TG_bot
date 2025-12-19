"""Document handlers for file uploads and processing.

Handles file uploads, processing, and analysis responses.
Supports multiple LLM providers with fallback.
"""

import logging
import tempfile
import uuid
from pathlib import Path

from aiogram import Router, F
from aiogram.types import Message, Document, File
from aiogram.fsm.context import FSMContext

from app.config import get_settings
from app.states.analysis import DocumentAnalysisStates
from app.services.file_processing.converter import FileConverter
from app.services.llm.llm_factory import LLMFactory
from app.utils.text_splitter import TextSplitter
from app.utils.cleanup import CleanupManager

logger = logging.getLogger(__name__)

router = Router()
config = get_settings()

# Initialize LLM factory
llm_factory = LLMFactory(
    primary_provider=config.LLM_PROVIDER,
    openai_api_key=config.OPENAI_API_KEY or None,
    openai_model=config.OPENAI_MODEL,
    replicate_api_token=config.REPLICATE_API_TOKEN or None,
    replicate_model=config.REPLICATE_MODEL,
)


@router.message(F.document)
async def handle_document(
    message: Message,
    state: FSMContext,
) -> None:
    """Handle document uploads.
    
    Args:
        message: User message with document
        state: FSM state
    """
    if not message.document:
        await message.answer("No document received.")
        return
    
    document: Document = message.document
    file_size = document.file_size or 0
    
    # Validate file size
    if file_size > config.MAX_FILE_SIZE:
        max_size_mb = config.MAX_FILE_SIZE / (1024 * 1024)
        await message.answer(
            f"⚠️ File is too large: {file_size / (1024 * 1024):.1f} MB\n"
            f"Maximum allowed: {max_size_mb:.1f} MB",
        )
        return
    
    # Set state
    await state.set_state(DocumentAnalysisStates.processing)
    
    # Show processing indicator
    processing_msg = await message.answer(
        "🔍 Processing your document...\n"
        "Downloading and extracting content..."
    )
    
    temp_user_dir = None
    files_to_cleanup: list[Path] = []
    
    try:
        # Create temp directory
        temp_base = Path(config.TEMP_DIR)
        temp_base.mkdir(exist_ok=True)
        temp_user_dir = CleanupManager.create_temp_directory(
            temp_base,
            message.from_user.id,
        )
        
        # Download file
        bot = message.bot
        file: File = await bot.get_file(document.file_id)
        
        if not file.file_path:
            await message.answer("❌ Failed to get file path.")
            return
        
        # Generate unique filename
        file_ext = Path(document.file_name or "document").suffix or ".bin"
        temp_file_path = temp_user_dir / f"{uuid.uuid4()}{file_ext}"
        files_to_cleanup.append(temp_file_path)
        
        await bot.download_file(file.file_path, temp_file_path)
        logger.info(f"Downloaded file: {temp_file_path} ({file_size} bytes)")
        
        # Extract text
        await processing_msg.edit_text(
            "🔍 Processing your document...\n"
            "Extracting text content..."
        )
        
        converter = FileConverter()
        extracted_text = converter.extract_text(temp_file_path, temp_user_dir)
        
        if not extracted_text or not extracted_text.strip():
            await message.answer(
                "⚠️ No text content found in the document. "
                "Please try another file."
            )
            await state.clear()
            return
        
        logger.info(f"Extracted {len(extracted_text)} characters")
        
        # Store in state for later use
        await state.update_data(
            extracted_text=extracted_text,
            original_filename=document.file_name,
            user_id=message.from_user.id,
        )
        
        # Analyze with AI
        await processing_msg.edit_text(
            "🔍 Processing your document...\n"
            f"🤖 Analyzing with {config.LLM_PROVIDER} AI..."
        )
        
        try:
            analysis_result = await llm_factory.analyze_document(
                extracted_text,
                "Analyze this document and provide key insights:",
                use_streaming=False,
            )
        
        except ValueError as e:
            logger.error(f"No LLM providers available: {e}")
            await message.answer(
                "❌ No LLM provider configured. "
                "Please set OPENAI_API_KEY or REPLICATE_API_TOKEN."
            )
            await state.clear()
            return
        
        except Exception as e:
            logger.error(f"LLM analysis error: {e}")
            # Try to use alternative provider if available
            available = llm_factory.get_available_providers()
            if len(available) > 1:
                other_provider = [
                    p for p in available if p != config.LLM_PROVIDER
                ][0]
                logger.info(f"Switching to {other_provider} provider")
                llm_factory.set_primary_provider(other_provider)
                try:
                    analysis_result = await llm_factory.analyze_document(
                        extracted_text,
                        "Analyze this document and provide key insights:",
                        use_streaming=False,
                    )
                except Exception as e2:
                    logger.error(f"Fallback also failed: {e2}")
                    await message.answer(f"❌ Analysis failed: {str(e2)}")
                    await state.clear()
                    return
            else:
                await message.answer(f"❌ Analysis failed: {str(e)}")
                await state.clear()
                return
        
        if not analysis_result:
            await message.answer(
                "❌ Analysis failed. Please try again later."
            )
            await state.clear()
            return
        
        logger.info(f"Analysis completed ({len(analysis_result)} chars)")
        
        # Send response
        await processing_msg.delete()
        
        splitter = TextSplitter()
        message_count = splitter.count_messages(analysis_result)
        
        if message_count <= 3:
            # Send as text messages
            chunks = splitter.split(analysis_result)
            for i, chunk in enumerate(chunks, 1):
                prefix = f"*[Part {i}/{len(chunks)}]*\n\n" if len(chunks) > 1 else ""
                await message.answer(
                    f"{prefix}{chunk}",
                    parse_mode="Markdown",
                )
        else:
            # Too long - send as file
            await send_analysis_as_file(
                message,
                analysis_result,
                document.file_name or "document",
                temp_user_dir,
                files_to_cleanup,
            )
        
        logger.info(
            f"Analysis sent to user {message.from_user.id} "
            f"({message_count} messages) using {config.LLM_PROVIDER}"
        )
        await state.clear()
    
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        await message.answer(
            f"❌ Error processing document: {str(e)}"
        )
        await state.clear()
    
    finally:
        # Cleanup
        if files_to_cleanup:
            await CleanupManager.cleanup_files_async(files_to_cleanup)
        if temp_user_dir:
            await CleanupManager.cleanup_directory_async(temp_user_dir)


async def send_analysis_as_file(
    message: Message,
    analysis: str,
    original_filename: str,
    temp_dir: Path,
    cleanup_list: list[Path],
) -> None:
    """Send analysis result as a Word document.
    
    Args:
        message: Telegram message to reply to
        analysis: Analysis text
        original_filename: Original file name
        temp_dir: Temporary directory
        cleanup_list: List to add file for cleanup
    """
    from docx import Document
    
    try:
        # Create Word document
        doc = Document()
        doc.add_heading(
            f"Analysis Report: {Path(original_filename).stem}",
            level=1,
        )
        doc.add_paragraph(analysis)
        
        # Save to temp file
        output_file = temp_dir / f"analysis_{uuid.uuid4()}.docx"
        doc.save(output_file)
        cleanup_list.append(output_file)
        
        # Send file
        await message.answer_document(
            open(output_file, "rb"),
            caption="📄 Analysis result (too long for messages, sent as document)",
        )
        
        logger.info(f"Analysis sent as file: {output_file}")
    
    except Exception as e:
        logger.error(f"Failed to send analysis as file: {e}")
        await message.answer(
            f"❌ Failed to send analysis file: {str(e)}"
        )
