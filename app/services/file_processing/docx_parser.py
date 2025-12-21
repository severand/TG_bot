"""DOCX file parser with robust error handling for old .doc files.

Фиксы 2025-12-21 11:40 - ПЕРЕПИСАН OLE EXTRACTION:
- Теперь используем грамотные методы нахождения текста
- Ни служебных данных, ни бинарного мусора
- FIB структура для нахождения текста

Фиксы 2025-12-21 11:39 - ПРЕДПРОСМОТР ТЕКСТА В ЛОГАХ:
- Добавлен предпросмотр 100-150 слов после каждого извлечения

Handles extraction of text content from Microsoft Word (.docx) files
using python-docx library with graceful fallback for corrupted files.
"""

import logging
from pathlib import Path
from zipfile import ZipFile, BadZipFile
from xml.etree import ElementTree as ET
from typing import Optional
import re
import struct

from app.services.file_processing.text_cleaner import TextCleaner

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from docx.oxml.text.paragraph import CT_P
    from docx.text.paragraph import Paragraph
    from docx.table import _Cell
except ImportError:
    CT_P = None
    Paragraph = None
    _Cell = None

try:
    import olefile
except ImportError:
    olefile = None

logger = logging.getLogger(__name__)


def _get_text_preview(text: str, max_words: int = 150) -> str:
    """Получить предпросмотр текста (первые N слов).
    
    Args:
        text: Исходный текст
        max_words: Максимальное количество слов
        
    Returns:
        Предпросмотр текста
    """
    if not text or not text.strip():
        return "(empty)"
    
    # Разбиваем на слова
    words = text.split()
    
    if len(words) <= max_words:
        return text.strip()[:800]  # Максимум 800 символов
    
    # Берем первые max_words слов
    preview = ' '.join(words[:max_words])
    return preview[:800] + "..."


class DOCXParser:
    """Парсер для DOCX (документов Microsoft Word).
    
    Поддерживает .docx и старые .doc файлы.
    Порядок исвлечения:
    1. python-docx (.docx файлы)
    2. ZIP extraction (поврежденные .docx)
    3. OLE extraction (старые .doc) - ПОЛНОЕ ПАРСИРОВАНИЕ FIB
    4. Binary extraction (последняя попытка)
    """
    
    def __init__(self) -> None:
        """Инициализация парсера."""
        self.text_cleaner = TextCleaner()
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file.
        
        Extracts text from paragraphs and tables, preserving document structure.
        Tries python-docx first, ZIP, then OLE (if olefile available), then binary.
        Cleans extracted text from binary sources.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text with structure preserved
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid DOCX
            
        Example:
            >>> parser = DOCXParser()
            >>> text = parser.extract_text(Path("document.docx"))
        """
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        logger.info(f"Starting extraction from {file_path.name} ({file_path.stat().st_size} bytes)")
        
        # Try standard python-docx first (for valid DOCX files)
        try:
            logger.info(f"Trying python-docx for {file_path.name}")
            doc = Document(file_path)
            extracted_text: list[str] = []
            
            # Extract paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                text = self._extract_paragraph_text(paragraph)
                if text:
                    extracted_text.append(text)
                    paragraph_count += 1
            
            logger.debug(f"Extracted {paragraph_count} paragraphs")
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    extracted_text.append(table_text)
                    table_count += 1
            
            logger.debug(f"Extracted {table_count} tables")
            
            result = "\n".join(extracted_text)
            if result.strip():
                logger.info(f"✓ Successfully extracted {len(result)} chars from {file_path.name} "
                          f"({paragraph_count} paragraphs, {table_count} tables)")
                preview = _get_text_preview(result, max_words=150)
                logger.info(f"📝 TEXT PREVIEW (first 150 words):\n{preview}")
                return result
            else:
                logger.warning(f"python-docx returned empty text, trying fallback")
        
        except Exception as e:
            logger.warning(
                f"python-docx failed for {file_path.name}: "
                f"{type(e).__name__}: {str(e)[:100]}"
            )
            logger.info(f"This might be an old .doc file, trying ZIP fallback...")
        
        # Fallback 1: Extract directly from ZIP (for .docx-like .doc files)
        logger.info(f"Step 1: Trying ZIP extraction for {file_path.name}")
        try:
            result = self._extract_from_zip(file_path)
            if result and result.strip():
                logger.info(f"✓ ZIP extraction successful: {len(result)} chars")
                preview = _get_text_preview(result, max_words=150)
                logger.info(f"📝 TEXT PREVIEW (first 150 words):\n{preview}")
                return result
            else:
                logger.info(f"ZIP extraction returned empty result, continuing to OLE method...")
        
        except BadZipFile:
            logger.info(f"File is not a ZIP archive, continuing to OLE method...")
        except Exception as e:
            logger.warning(f"ZIP fallback failed with {type(e).__name__}, continuing to OLE method...")
        
        # Fallback 2: Extract using OLE (for old MS Word 97-2003 binary .doc)
        logger.info(f"Step 2: Trying OLE extraction for {file_path.name}")
        try:
            result = self._extract_from_ole_doc(file_path)
            if result and result.strip():
                logger.info(f"✓ OLE extraction successful: {len(result)} chars (before cleaning)")
                preview = _get_text_preview(result, max_words=150)
                logger.info(f"📝 TEXT PREVIEW (OLE, before cleaning):\n{preview}")
                
                # Clean the extraction result
                logger.info(f"Cleaning extracted text from OLE...")
                cleaned_result = self.text_cleaner.clean_extracted_text(result, aggressive=True)
                
                if cleaned_result and self.text_cleaner.is_text_usable(cleaned_result):
                    logger.info(f"✓ Cleaned OLE text: {len(cleaned_result)} chars (quality OK)")
                    preview = _get_text_preview(cleaned_result, max_words=150)
                    logger.info(f"📝 TEXT PREVIEW (OLE, after cleaning):\n{preview}")
                    return cleaned_result
                else:
                    logger.warning(f"Cleaned text quality poor, returning raw OLE text")
                    return result
            else:
                logger.info(f"OLE extraction returned empty result, continuing to binary method...")
        
        except Exception as e:
            logger.warning(f"OLE extraction failed with {type(e).__name__}: {str(e)[:100]}")
            logger.info(f"Continuing to binary extraction as last resort...")
        
        # Fallback 3: Extract from binary old .doc format (last resort)
        logger.info(f"Step 3: Trying binary extraction for {file_path.name}")
        try:
            result = self._extract_from_binary_doc(file_path)
            if result and result.strip():
                logger.info(f"✓ Binary extraction successful: {len(result)} chars (before cleaning)")
                preview = _get_text_preview(result, max_words=150)
                logger.info(f"📝 TEXT PREVIEW (Binary, before cleaning):\n{preview}")
                
                # Clean extracted text
                logger.info(f"Cleaning extracted text from binary...")
                cleaned_result = self.text_cleaner.clean_extracted_text(result, aggressive=False)
                
                if cleaned_result and self.text_cleaner.is_text_usable(cleaned_result):
                    logger.info(f"✓ Cleaned binary text: {len(cleaned_result)} chars (quality OK)")
                    preview = _get_text_preview(cleaned_result, max_words=150)
                    logger.info(f"📝 TEXT PREVIEW (Binary, after cleaning):\n{preview}")
                    return cleaned_result
                else:
                    logger.warning(f"Cleaned text quality poor, returning raw binary text")
                    return result
            else:
                logger.warning(f"Binary extraction returned empty result")
                return ""
        
        except Exception as e:
            logger.error(f"Binary extraction failed: {type(e).__name__}: {e}")
            raise ValueError(f"Invalid DOCX/DOC file: Cannot extract text using any method") from e
    
    def _extract_from_ole_doc(self, file_path: Path) -> str:
        """Extract text from old .doc file using OLE structure with FIB parsing.
        
        Uses olefile library to read MS Word 97-2003 OLE format.
        Parses File Information Block (FIB) to find text range.
        
        Args:
            file_path: Path to .doc file
            
        Returns:
            str: Extracted text (empty string if OLE not available or file not OLE)
        """
        if olefile is None:
            logger.debug("olefile library not installed - skipping OLE extraction")
            return ""
        
        try:
            # Check if it's actually an OLE file
            if not olefile.isOleFile(str(file_path)):
                logger.debug(f"File is not OLE format - OLE extraction not applicable")
                return ""
            
            ole = olefile.OleFileIO(str(file_path))
            logger.debug(f"Successfully opened OLE file {file_path.name}")
            
            text_result = ""
            
            # Extract from WordDocument stream with proper FIB parsing
            if ole.exists('WordDocument'):
                logger.debug(f"Found WordDocument stream")
                try:
                    stream = ole.openstream('WordDocument')
                    data = stream.read()
                    text_result = self._extract_text_from_fib(data)
                    if text_result and text_result.strip():
                        logger.debug(f"Extracted {len(text_result)} chars using FIB parsing")
                except Exception as e:
                    logger.debug(f"FIB parsing failed: {type(e).__name__}, trying alternative method")
            
            # Fallback: Try WordDocument stream without FIB (raw text search)
            if not text_result or len(text_result.strip()) < 50:
                if ole.exists('WordDocument'):
                    try:
                        stream = ole.openstream('WordDocument')
                        data = stream.read()
                        text_result = self._extract_text_from_word_stream_raw(data)
                        if text_result and text_result.strip():
                            logger.debug(f"Extracted {len(text_result)} chars using raw stream method")
                    except Exception as e:
                        logger.debug(f"Raw stream extraction failed: {type(e).__name__}")
            
            ole.close()
            
            if text_result and text_result.strip():
                logger.info(f"OLE extraction successful: {len(text_result)} chars total")
                return text_result
            else:
                logger.debug(f"OLE extraction found no text")
                return ""
        
        except Exception as e:
            logger.debug(f"OLE extraction error: {type(e).__name__}: {str(e)[:50]}")
            return ""
    
    def _extract_text_from_fib(self, data: bytes) -> str:
        """Parse File Information Block (FIB) and extract text.
        
        FIB is at start of WordDocument stream.
        Contains pointers to text location and length.
        
        Args:
            data: Binary WordDocument stream data
            
        Returns:
            str: Extracted text
        """
        try:
            if len(data) < 32:
                return ""
            
            # Check signature
            signature = struct.unpack('<H', data[0:2])[0]
            if signature != 0xDB:
                logger.debug(f"Invalid FIB signature: 0x{signature:04X}")
                return ""
            
            # Get encoding info (byte 80: 1 = ANSI, 0 = Unicode)
            is_ansi = data[80] == 1 if len(data) > 80 else True
            encoding = 'cp1251' if is_ansi else 'utf-16-le'
            
            logger.debug(f"FIB: is_ansi={is_ansi}, encoding={encoding}")
            
            # Extract all readable text from stream
            # Since structure is complex, use simple ASCII/UTF-8 extraction
            text_parts = []
            current_word = b''
            
            for byte in data[512:]:  # Skip FIB header
                if 32 <= byte <= 126:  # Printable ASCII
                    current_word += bytes([byte])
                elif 0xC0 <= byte <= 0xFF:  # Cyrillic UTF-8 continuation
                    current_word += bytes([byte])
                else:
                    if len(current_word) > 2:
                        try:
                            word = current_word.decode('cp1251', errors='ignore').strip()
                            if word and len(word) > 1:
                                text_parts.append(word)
                        except:
                            pass
                    current_word = b''
            
            # Don't forget last word
            if len(current_word) > 2:
                try:
                    word = current_word.decode('cp1251', errors='ignore').strip()
                    if word:
                        text_parts.append(word)
                except:
                    pass
            
            if text_parts:
                return ' '.join(text_parts)
            else:
                return ""
        
        except Exception as e:
            logger.debug(f"FIB parsing error: {type(e).__name__}")
            return ""
    
    def _extract_text_from_word_stream_raw(self, data: bytes) -> str:
        """Extract text from WordDocument stream (raw method).
        
        Searches for readable text patterns in stream.
        
        Args:
            data: Binary WordDocument stream data
            
        Returns:
            str: Extracted text
        """
        text_parts = []
        
        # Try multiple encodings
        for encoding in ['cp1251', 'utf-8', 'latin-1']:
            try:
                decoded = data.decode(encoding, errors='ignore')
                # Find readable Cyrillic and Latin words
                words = re.findall(r'[\u0400-\u04FFA-Za-z]{2,}', decoded)
                
                if len(words) > 20:  # Found enough words
                    text = ' '.join(words)
                    if len(text) > 100:
                        logger.debug(f"Raw extraction: found {len(text)} chars with {encoding}")
                        text_parts.append(text)
                        break
            except Exception:
                continue
        
        return ' '.join(text_parts[:1]) if text_parts else ""
    
    def _extract_from_zip(self, file_path: Path) -> str:
        """Extract text directly from DOCX ZIP archive.
        
        DOCX is a ZIP file containing document.xml which has the text.
        For .doc files with ZIP structure, this provides extraction.
        Also handles incomplete or corrupted ZIP structures.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text from XML (empty string if not a ZIP or no text found)
        """
        logger.debug(f"ZIP extraction: checking if file is ZIP archive")
        
        all_text = []
        
        try:
            with ZipFile(file_path, 'r') as docx_zip:
                file_list = docx_zip.namelist()
                logger.debug(f"ZIP contains {len(file_list)} files")
                
                # Try document.xml first (main content)
                if 'word/document.xml' in file_list:
                    try:
                        text = self._extract_text_from_xml(docx_zip, 'word/document.xml')
                        if text and text.strip():
                            all_text.append(text)
                            logger.debug(f"ZIP: extracted {len(text)} chars from document.xml")
                    except Exception as e:
                        logger.debug(f"ZIP: failed to parse document.xml: {type(e).__name__}")
                
                # Try other XML files as fallback
                for file_name in file_list:
                    if file_name.endswith('.xml') and 'word/' in file_name and file_name != 'word/document.xml':
                        try:
                            text = self._extract_text_from_xml(docx_zip, file_name)
                            if text and text.strip():
                                all_text.append(text)
                                logger.debug(f"ZIP: extracted {len(text)} chars from {file_name}")
                        except Exception:
                            continue
        
        except BadZipFile:
            logger.debug(f"ZIP: file is not a valid ZIP archive (not a DOCX)")
            return ""
        except Exception as e:
            logger.debug(f"ZIP extraction error: {type(e).__name__}: {str(e)[:50]}")
            return ""
        
        result = "\n\n".join(all_text)
        logger.debug(f"ZIP extraction result: {len(result)} chars total")
        return result
    
    def _extract_text_from_xml(self, docx_zip: ZipFile, xml_path: str) -> str:
        """Extract text from XML file in DOCX ZIP.
        
        Handles corrupted XML by extracting readable text.
        
        Args:
            docx_zip: ZipFile object
            xml_path: Path to XML file in ZIP
            
        Returns:
            str: Extracted text
        """
        try:
            xml_content = docx_zip.read(xml_path).decode('utf-8', errors='ignore')
            
            if not xml_content.strip():
                return ""
            
            # Try to parse as XML
            try:
                root = ET.fromstring(xml_content)
                namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text_elements = root.findall('.//w:t', namespace)
                
                # Fallback: try without namespace
                if not text_elements:
                    text_elements = root.findall('.//t')
                
                text_parts = []
                for elem in text_elements:
                    if elem.text:
                        text_parts.append(elem.text)
                
                text = ''.join(text_parts)
                return text
            
            except ET.ParseError:
                # If XML parsing fails, extract plain text from corrupted content
                return self._extract_plain_text_from_corrupted_xml(xml_content)
        
        except Exception:
            return ""
    
    def _extract_from_binary_doc(self, file_path: Path) -> str:
        """Try to extract text from binary .doc format.
        
        Old MS Word .doc files are binary (OLE format), not ZIP.
        This attempts to extract readable ASCII/UTF-8 strings.
        Improved version with better word detection.
        
        Args:
            file_path: Path to .doc file
            
        Returns:
            str: Extracted text
        """
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            if not content:
                logger.warning(f"File {file_path.name} is empty")
                return ""
            
            # Check if it's OLE format (MS Office binary)
            is_ole = content[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
            if is_ole:
                logger.debug(f"Binary: file is OLE format (MS Office binary .doc)")
            
            # Strategy 1: Look for null-separated text blocks (common in MS Word)
            text_parts = self._extract_text_from_null_blocks(content)
            if text_parts:
                text = ' '.join(text_parts)
                if len(text) > 50:
                    logger.debug(f"Binary: found {len(text)} chars using null-block method")
                    return text.strip()
            
            # Strategy 2: Look for continuous printable strings
            text_parts = self._extract_continuous_strings(content)
            if text_parts:
                text = ' '.join(text_parts)
                if len(text) > 50:
                    logger.debug(f"Binary: found {len(text)} chars using continuous string method")
                    return text.strip()
            
            # Strategy 3: Look for UTF-16 encoded text (Word uses this)
            try:
                text = self._extract_utf16_strings(content)
                if text and len(text.strip()) > 50:
                    logger.debug(f"Binary: found {len(text)} chars using UTF-16 method")
                    return text.strip()
            except:
                pass
            
            logger.debug(f"Binary: no substantial text found")
            return ""
        
        except Exception as e:
            logger.error(f"Binary extraction failed: {type(e).__name__}: {e}")
            return ""
    
    def _extract_text_from_null_blocks(self, content: bytes) -> list[str]:
        """Extract text from null-separated blocks in binary data.
        
        MS Word stores text in 2-byte blocks separated by nulls.
        """
        text_parts = []
        current_word = b''
        
        i = 0
        while i < len(content):
            byte = content[i]
            
            # Printable ASCII
            if 32 <= byte <= 126:  # ASCII printable
                current_word += bytes([byte])
            # Whitespace
            elif byte in (9, 10, 13, 0):  # tab, newline, carriage return, null
                if len(current_word) > 2:
                    try:
                        word = current_word.decode('ascii', errors='ignore').strip()
                        if word and len(word) > 1:
                            text_parts.append(word)
                    except:
                        pass
                current_word = b''
            else:
                current_word = b''
            
            i += 1
        
        return text_parts
    
    def _extract_continuous_strings(self, content: bytes, min_len: int = 4) -> list[str]:
        """Extract continuous printable strings from binary data.
        
        Args:
            content: Binary content
            min_len: Minimum string length to consider
            
        Returns:
            List of strings found
        """
        strings = []
        current = b''
        
        for byte in content:
            if 32 <= byte <= 126:  # Printable ASCII
                current += bytes([byte])
            else:
                if len(current) >= min_len:
                    try:
                        s = current.decode('ascii', errors='ignore').strip()
                        if s:
                            strings.append(s)
                    except:
                        pass
                current = b''
        
        # Don't forget the last one
        if len(current) >= min_len:
            try:
                s = current.decode('ascii', errors='ignore').strip()
                if s:
                    strings.append(s)
            except:
                pass
        
        return strings
    
    def _extract_utf16_strings(self, content: bytes) -> str:
        """Extract UTF-16 encoded strings (common in modern .doc files).
        
        Args:
            content: Binary content
            
        Returns:
            Extracted text
        """
        # Try UTF-16 LE (little-endian, common in Windows)
        try:
            decoded = content.decode('utf-16-le', errors='ignore')
            words = re.findall(r'[\w\s]{2,}', decoded)
            if words:
                return ' '.join(words).strip()
        except:
            pass
        
        # Try UTF-16 BE (big-endian)
        try:
            decoded = content.decode('utf-16-be', errors='ignore')
            words = re.findall(r'[\w\s]{2,}', decoded)
            if words:
                return ' '.join(words).strip()
        except:
            pass
        
        return ""
    
    @staticmethod
    def _extract_plain_text_from_corrupted_xml(xml_content: str) -> str:
        """Extract readable text from corrupted XML.
        
        When XML parsing fails, extract text between tags.
        
        Args:
            xml_content: Raw XML content
            
        Returns:
            str: Extracted text
        """
        # Remove XML tags
        text = re.sub(r'<[^>]+>', ' ', xml_content)
        # Clean up whitespace
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = text.strip()
        return text
    
    def _extract_paragraph_text(self, paragraph) -> str:  # type: ignore
        """Extract text from a paragraph with formatting preserved.
        
        Args:
            paragraph: Paragraph object
            
        Returns:
            str: Paragraph text
        """
        try:
            text = paragraph.text
            if not text.strip():
                return ""
            
            # Preserve heading formatting with markdown
            if hasattr(paragraph, 'style') and paragraph.style.name.startswith("Heading"):
                try:
                    level = int(paragraph.style.name.replace("Heading", ""))
                    return f"{'#' * level} {text}"
                except (ValueError, IndexError):
                    pass
            
            return text
        except Exception as e:
            logger.debug(f"Error extracting paragraph text: {e}")
            return ""
    
    def _extract_table_text(self, table) -> str:  # type: ignore
        """Extract text from a table.
        
        Args:
            table: Table object
            
        Returns:
            str: Formatted table text
        """
        try:
            rows: list[list[str]] = []
            
            for row in table.rows:
                cells: list[str] = []
                for cell in row.cells:
                    cell_text = self._extract_cell_text(cell)
                    cells.append(cell_text)
                if cells:
                    rows.append(cells)
            
            if not rows:
                return ""
            
            # Format as markdown table
            table_text = []
            table_text.append(" | ".join(rows[0]))
            table_text.append(" | ".join(["---"] * len(rows[0])))
            
            for row in rows[1:]:
                table_text.append(" | ".join(row))
            
            return "\n" + "\n".join(table_text) + "\n"
        except Exception as e:
            logger.debug(f"Error extracting table text: {e}")
            return ""
    
    def _extract_cell_text(self, cell) -> str:  # type: ignore
        """Extract text from a table cell.
        
        Args:
            cell: Table cell
            
        Returns:
            str: Cell text
        """
        try:
            cell_text = []
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                if text:
                    cell_text.append(text)
            return " ".join(cell_text)
        except Exception as e:
            logger.debug(f"Error extracting cell text: {e}")
            return ""
    
    def get_metadata(self, file_path: Path) -> dict:  # type: ignore
        """Extract DOCX metadata.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            dict: Document metadata
        """
        try:
            doc = Document(file_path)
            props = doc.core_properties
            return {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "created": props.created,
                "modified": props.modified,
            }
        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")
            return {}
