"""Specialized parser for old MS Word .doc binary format.

Uses LibreOffice (soffice) to convert .doc to .docx on the fly.
LibreOffice works perfectly with old Russian .doc files (CP1251).

Requirements:
- LibreOffice installed and in PATH
- soffice command available

Windows:
- Install LibreOffice: https://www.libreoffice.org/
- Add "C:\\Program Files\\LibreOffice\\program" to PATH
"""

import logging
import subprocess
import shutil
import uuid
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DOCParser:
    """Specialized parser for old binary .doc format using LibreOffice.
    
    Old MS Word .doc files are binary and hard to parse.
    LibreOffice converts them to .docx perfectly.
    """
    
    def __init__(self) -> None:
        self._soffice_path = self._find_libreoffice()
        if self._soffice_path:
            logger.info(f"✓ LibreOffice found at: {self._soffice_path}")
        else:
            logger.warning("⚠ LibreOffice (soffice) not found!")
    
    @staticmethod
    def _find_libreoffice() -> Optional[str]:
        """Find path to LibreOffice executable."""
        # Check in PATH
        if shutil.which("soffice"):
            return "soffice"
        
        # Check common Windows paths
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in paths:
            if Path(path).exists():
                return path
        
        return None
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from .doc file by converting to .docx with LibreOffice.
        
        Args:
            file_path: Path to .doc file
            
        Returns:
            str: Extracted text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If LibreOffice failed
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not self._soffice_path:
            raise ValueError(
                "LibreOffice not found! Please install LibreOffice and add to PATH. "
                "Download: https://www.libreoffice.org/"
            )
        
        logger.info(f"Converting {file_path.name} to .docx using LibreOffice...")
        
        # Create temp output dir in the same folder
        output_dir = file_path.parent
        
        try:
            # Convert .doc to .docx
            # soffice --headless --convert-to docx "file.doc" --outdir "output_dir"
            cmd = [
                self._soffice_path,
                "--headless",
                "--convert-to",
                "docx",
                str(file_path),
                "--outdir",
                str(output_dir)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")
                raise ValueError(f"LibreOffice conversion failed: {result.stderr}")
            
            # Find the converted file (same name but .docx)
            docx_path = output_dir / (file_path.stem + ".docx")
            
            if not docx_path.exists():
                raise ValueError(f"Converted .docx not found at {docx_path}")
            
            logger.info(f"✓ Conversion successful: {docx_path.name}")
            
            # Now extract text from the new .docx using python-docx
            # We import here to avoid circular dependency
            from docx import Document
            
            doc = Document(docx_path)
            extracted_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        extracted_text.append(" | ".join(row_text))
            
            text = "\n".join(extracted_text)
            
            # Clean up the temporary .docx file
            try:
                docx_path.unlink()
            except Exception:
                pass
            
            if text.strip():
                return text
            else:
                raise ValueError("Extracted text is empty")

        except Exception as e:
            logger.error(f"Error processing .doc file: {e}")
            raise ValueError(f"Failed to process .doc file: {e}")
